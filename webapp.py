"""Web interface for Carbon Footprint Optimization Engine (CfoE)."""

from __future__ import annotations

import asyncio
from queue import Queue
import csv
import json
import os
import threading
from textwrap import wrap
from datetime import datetime, timezone
from functools import lru_cache
from pathlib import Path
from typing import Any, Dict, List, Optional
from uuid import uuid4

from dotenv import load_dotenv
from fastapi import FastAPI, Header, HTTPException, Request, WebSocket, WebSocketDisconnect, BackgroundTasks
from fastapi.responses import FileResponse, JSONResponse
from fastapi.staticfiles import StaticFiles
from pydantic import BaseModel, Field
from docx import Document
from reportlab.lib.pagesizes import A4

from agents.calculation_agent import calculate_carbon_score
from agents.policy_agent import enforce_policy_hitl
from agents.registry_agent import validate_registry_id, get_entity_info
from agents.trajectory_agent import calculate_trajectory, check_compliance_trajectory
from agents.credit_agent import calculate_carbon_credits, get_supplier_credits, get_supplier_credit_history, get_leaderboard
from orchestrators.root_coordinator import create_root_coordinator

from config.groq_config import get_groq_client
from blockchain_client import get_blockchain_client
from carbon_token_manager import get_token_manager

try:
    from groq import Groq
except ImportError:  # pragma: no cover
    Groq = None

# ── X402 Payment constants ─────────────────────────────────────────────────
# Shared secret that lets the internal simulator bypass the payment gate.
X402_INTERNAL_SECRET = os.getenv("X402_INTERNAL_SECRET", "cfoe-internal-bypass-secret")
# Required ALGO per external audit call
AUDIT_PAYMENT_ALGO = 0.05
# Required ALGO per report access
REPORT_PAYMENT_ALGO = 0.02

BASE_DIR = Path(__file__).resolve().parent
WEB_DIR = BASE_DIR / "web"
STATIC_DIR = WEB_DIR / "static"
DATA_DIR = BASE_DIR / "data"
HISTORY_PATH = DATA_DIR / "audit_history.json"
PENDING_PATH = DATA_DIR / "pending_approvals.json"
OUTPUT_DIR = BASE_DIR / "outputs"
OUTPUT_CSV_PATH = OUTPUT_DIR / "audits_master.csv"

lock = threading.Lock()
log_queue = Queue()
active_websockets: List[WebSocket] = []

def broadcast_log_sync(log_msg: Dict[str, Any]) -> None:
    """Put log in queue and try to send to active websockets"""
    log_queue.put(log_msg)
    # Don't create async tasks here - let the websocket endpoint handle it
    # The log queue will be consumed by the websocket endpoint


class AuditRequest(BaseModel):
    supplier_name: str = Field(min_length=2, max_length=120)
    emissions: float = Field(ge=0)
    violations: int = Field(ge=0, le=500)
    notes: str = Field(default="", max_length=1000)
    sector: str = Field(default="default")  # Phase 1: Sector-specific targets
    production_volume: float = Field(default=None, ge=0)  # Phase 2: Normalized metrics
    production_unit: str = Field(default="tonne")  # Phase 2: Unit for normalization
    registry_id: str = Field(default="")  # Phase 3: Entity registry
    baseline_year: int = Field(default=2023)  # Phase 1: Pro-rata calculation
    target_year: int = Field(default=2027)  # Phase 1: Pro-rata calculation


class AuditResponse(BaseModel):
    job_id: str
    audit_id: str
    timestamp: str
    supplier_name: str
    emissions: float
    violations: int
    notes: str
    risk_score: float
    classification: str
    emissions_score: float
    violations_score: float
    external_risk_score: float
    policy_decision: str
    human_approval_required: bool
    policy_reason: str
    recommended_action: str
    report_text: str
    report_source: str
    download_links: Dict[str, str]
    status: str  # "completed" or "pending_approval"


class ApprovalRequest(BaseModel):
    audit_id: str
    decision: str  # "approve" or "reject"
    approver_name: str = Field(min_length=2, max_length=100)
    approval_notes: str = Field(default="", max_length=500)


app = FastAPI(
    title="CfoE Dashboard API",
    description="Interactive interface for supplier ESG risk audits.",
    version="1.0.0",
)


@app.middleware("http")
async def add_no_cache_headers(request: Request, call_next):
    response = await call_next(request)

    # Prevent stale frontend assets during local development.
    if request.url.path == "/" or request.url.path.startswith("/static/"):
        response.headers["Cache-Control"] = "no-store, no-cache, must-revalidate, max-age=0"
        response.headers["Pragma"] = "no-cache"
        response.headers["Expires"] = "0"

    return response

# Ensure static output directory exists before mount.
OUTPUT_DIR.mkdir(parents=True, exist_ok=True)

# Mount simulator directory
SIMULATOR_DIR = BASE_DIR / "simulator"
app.mount("/static", StaticFiles(directory=str(STATIC_DIR)), name="static")
app.mount("/outputs", StaticFiles(directory=str(OUTPUT_DIR)), name="outputs")


@lru_cache(maxsize=1)
def get_client():
    load_dotenv()
    api_key = __import__("os").getenv("GROQ_API_KEY")
    if not api_key or Groq is None:
        return None
    try:
        return Groq(api_key=api_key)
    except Exception:
        return None


def ensure_storage() -> None:
    DATA_DIR.mkdir(parents=True, exist_ok=True)
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)

    if not HISTORY_PATH.exists():
        HISTORY_PATH.write_text("[]", encoding="utf-8")
    
    if not PENDING_PATH.exists():
        PENDING_PATH.write_text("[]", encoding="utf-8")

    # Ensure credit ledger files exist (used by credit_agent)
    credit_ledger_path = DATA_DIR / "credit_ledger.json"
    if not credit_ledger_path.exists():
        credit_ledger_path.write_text("{}", encoding="utf-8")

    # Ensure credit issuance ledger exists (used by carbon_token_manager)
    credit_issuance_path = DATA_DIR / "credit_issuance_ledger.json"
    if not credit_issuance_path.exists():
        credit_issuance_path.write_text("[]", encoding="utf-8")
    else:
        content = credit_issuance_path.read_text(encoding="utf-8").strip()
        if not content:
            credit_issuance_path.write_text("[]", encoding="utf-8")

    # Ensure token state file exists (used by carbon_token_manager)
    token_state_path = DATA_DIR / "token_state.json"
    if not token_state_path.exists():
        token_state_path.write_text("{}", encoding="utf-8")
    else:
        content = token_state_path.read_text(encoding="utf-8").strip()
        if not content:
            token_state_path.write_text("{}", encoding="utf-8")

    # Ensure blockchain ledger exists (used by blockchain_client for persistence)
    blockchain_ledger_path = DATA_DIR / "blockchain_ledger.json"
    if not blockchain_ledger_path.exists():
        blockchain_ledger_path.write_text(
            '{"score_anchors": [], "hitl_decisions": [], "report_hashes": []}',
            encoding="utf-8",
        )

    # ── New data files for ASA economy (Parts 2–5) ───────────────────────
    # pending_mints.json — CCC tokens queued for walletless suppliers
    pending_mints_path = DATA_DIR / "pending_mints.json"
    if not pending_mints_path.exists():
        pending_mints_path.write_text("[]", encoding="utf-8")

    # compliance_bonds.json — active / historic compliance bonds
    compliance_bonds_path = DATA_DIR / "compliance_bonds.json"
    if not compliance_bonds_path.exists():
        compliance_bonds_path.write_text("{}", encoding="utf-8")

    # marketplace_listings.json — P2P credit listings
    marketplace_path = DATA_DIR / "marketplace_listings.json"
    if not marketplace_path.exists():
        marketplace_path.write_text('{"listings": {}, "counter": 0}', encoding="utf-8")

    # staking_positions.json — CCC staking positions
    staking_path = DATA_DIR / "staking_positions.json"
    if not staking_path.exists():
        staking_path.write_text("{}", encoding="utf-8")

    # ── X402 Agentic Commerce data files ──────────────────────────────────
    # agent_wallets.json — agent wallet addresses (no private keys)
    agent_wallets_path = DATA_DIR / "agent_wallets.json"
    if not agent_wallets_path.exists():
        agent_wallets_path.write_text("{}", encoding="utf-8")

    # agent_payments.json — X402 payment ledger
    agent_payments_path = DATA_DIR / "agent_payments.json"
    if not agent_payments_path.exists():
        agent_payments_path.write_text("[]", encoding="utf-8")

    # encrypted_reports.json — encrypted audit report registry
    encrypted_reports_path = DATA_DIR / "encrypted_reports.json"
    if not encrypted_reports_path.exists():
        encrypted_reports_path.write_text("{}", encoding="utf-8")

    if not OUTPUT_CSV_PATH.exists():
        with OUTPUT_CSV_PATH.open("w", encoding="utf-8", newline="") as csv_file:
            writer = csv.DictWriter(
                csv_file,
                fieldnames=[
                    "audit_id",
                    "timestamp",
                    "supplier_name",
                    "emissions",
                    "violations",
                    "risk_score",
                    "classification",
                    "policy_decision",
                    "human_approval_required",
                    "report_source",
                ],
            )
            writer.writeheader()


def load_history() -> List[Dict[str, Any]]:
    ensure_storage()
    with lock:
        try:
            content = HISTORY_PATH.read_text(encoding="utf-8").strip() or "[]"
            data = json.loads(content)
            return data if isinstance(data, list) else []
        except json.JSONDecodeError:
            return []


def save_history(history: List[Dict[str, Any]]) -> None:
    with lock:
        HISTORY_PATH.write_text(json.dumps(history, indent=2), encoding="utf-8")


def load_pending() -> List[Dict[str, Any]]:
    ensure_storage()
    with lock:
        try:
            content = PENDING_PATH.read_text(encoding="utf-8").strip() or "[]"
            data = json.loads(content)
            return data if isinstance(data, list) else []
        except json.JSONDecodeError:
            return []


def save_pending(pending: List[Dict[str, Any]]) -> None:
    with lock:
        PENDING_PATH.write_text(json.dumps(pending, indent=2), encoding="utf-8")


def make_audit_prompt(req: AuditRequest) -> str:
    return f"""
Conduct a comprehensive ESG audit for the following supplier:

Supplier Name: {req.supplier_name}
Annual CO2 Emissions: {req.emissions} tons
Regulatory Violations: {req.violations}
Additional Notes: {req.notes or 'N/A'}

Please provide a complete risk assessment and recommendations.
"""


def build_fallback_report(req: AuditRequest, risk_data: Dict[str, Any], policy_data: Dict[str, Any], blockchain_data: Optional[Dict[str, Any]] = None, credit_data: Optional[Dict[str, Any]] = None) -> str:
    report = (
        "Executive Summary\n"
        f"Supplier: {req.supplier_name}\n"
        f"Risk Score: {risk_data['risk_score']} ({risk_data['classification']})\n\n"
        "Key Findings\n"
        f"- Emissions contribution score: {risk_data['emissions_score']}\n"
        f"- Violations contribution score: {risk_data['violations_score']}\n"
        f"- Policy decision: {policy_data['decision']}\n\n"
        "Recommended Action\n"
        f"{policy_data['recommended_action']}\n"
    )
    
    # carbon credits
    if credit_data:
        report += "\n\n" + "="*60 + "\n"
        report += "CARBON CREDITS AWARDED\n"
        report += "="*60 + "\n\n"
        report += f"Base Credits: {credit_data.get('credits_earned', 0)}\n"
        if credit_data.get('streak_bonus', 0) > 0:
            report += f"Streak Bonus: +{credit_data['streak_bonus']}\n"
        if credit_data.get('improvement_bonus', 0) > 0:
            report += f"Improvement Bonus: +{credit_data['improvement_bonus']}\n"
        report += f"Total Earned: {credit_data.get('total_credits', 0)}\n"
        report += f"New Balance: {credit_data.get('new_total', 0)} credits\n"
        if credit_data.get('badges_earned'):
            report += f"Badges: {', '.join(credit_data['badges_earned'])}\n"
    
    # blockchain verification 
    if blockchain_data:
        report += "\n\n" + "="*60 + "\n"
        report += "BLOCKCHAIN VERIFICATION\n"
        report += "="*60 + "\n\n"
        
        if blockchain_data.get("score_tx"):
            report += f"Score Anchor TX: {blockchain_data['score_tx']}\n"
        if blockchain_data.get("score_hash"):
            report += f"Input Data Hash: {blockchain_data['score_hash']}\n"
        if blockchain_data.get("report_hash"):
            report += f"Report SHA-256: {blockchain_data['report_hash']}\n"
        if blockchain_data.get("verification_code"):
            report += f"Verification Code: {blockchain_data['verification_code']}\n"
        if blockchain_data.get("report_tx"):
            report += f"Report Hash TX: {blockchain_data['report_tx']}\n"
        if blockchain_data.get("credit_tx"):
            report += f"Carbon Credits TX: {blockchain_data['credit_tx']}\n"
        if blockchain_data.get("hitl_tx"):
            report += f"HITL Decision TX: {blockchain_data['hitl_tx']}\n"
        
        report += f"\nOn-Chain Status: {'✓ Verified' if blockchain_data.get('on_chain') else '✗ Local Only'}\n"
        report += "\nTo verify this report:\n"
        report += "1. Calculate SHA-256 hash of this report text\n"
        report += "2. Compare with the Report SHA-256 above\n"
        report += "3. Check transactions on Algorand Explorer:\n"
        report += "   https://testnet.algoexplorer.io/tx/[TX_ID]\n"
    
    return report


def run_audit(req: AuditRequest) -> Dict[str, Any]:
    broadcast_log_sync({"type": "info", "message": f"Starting audit for {req.supplier_name}..."})
    
    # Check if wallet is connected
    bc = get_blockchain_client()
    if not bc.wallet_connected:
        broadcast_log_sync({"type": "warning", "message": "⚠ Wallet not connected - transactions will be stored locally"})
    
    broadcast_log_sync({"type": "info", "message": "[1/6] Calculating ESG risk scores..."})
    
    # Calculate audit date for pro-rata
    from datetime import datetime
    audit_date = datetime.now()
    
    risk_data = calculate_carbon_score(
        emissions=req.emissions, 
        violations=req.violations,
        sector=req.sector,
        production_volume=req.production_volume,
        audit_date=audit_date
    )
    
    broadcast_log_sync({"type": "success", "message": f"✓ Risk Score: {risk_data['risk_score']} ({risk_data['classification']}) - Sector: {risk_data['sector']}"})
    
    # Blockchain: Anchor score
    bc = get_blockchain_client()
    if bc.wallet_connected and bc.address:
        broadcast_log_sync({"type": "info", "message": f"[Blockchain] Using wallet: {bc.address[:16]}..."})
    else:
        broadcast_log_sync({"type": "info", "message": "[Blockchain] Wallet not connected - storing locally"})
    
    try:
        score_anchor = bc.anchor_score(
            supplier_name=req.supplier_name,
            risk_score=risk_data["risk_score"],
            classification=risk_data["classification"],
            emissions=req.emissions,
            violations=req.violations,
            emissions_score=risk_data["emissions_score"],
            violations_score=risk_data["violations_score"],
            external_risk_score=risk_data.get("external_risk_score", 0.0)
        )
    except Exception as e:
        broadcast_log_sync({"type": "error", "message": f"✗ Score anchor exception: {str(e)}"})
        score_anchor = None
    
    if score_anchor:
        tx_id = score_anchor.get('tx_id') or score_anchor.get('local_id', 'LOCAL')
        broadcast_log_sync({"type": "success", "message": f"✓ Score anchored on blockchain: {tx_id[:20]}..."})
    else:
        broadcast_log_sync({"type": "error", "message": "✗ Score anchor failed - returned None"})
        score_anchor = {"local_id": "FALLBACK", "on_chain": False}
    
    broadcast_log_sync({"type": "info", "message": "[2/6] Enforcing policy rules..."})
    policy_data = enforce_policy_hitl(risk_score=risk_data["risk_score"], supplier_name=req.supplier_name)
    broadcast_log_sync({"type": "success", "message": f"✓ Policy Decision: {policy_data['decision']}"})

    # Prepare blockchain data for report
    blockchain_data_for_report = {
        "score_tx": score_anchor.get("tx_id") or score_anchor.get("local_id") if score_anchor else None,
        "score_hash": score_anchor.get("data_hash") if score_anchor else None,
        "on_chain": score_anchor.get("on_chain", False) if score_anchor else False
    }

    report_source = "deterministic-fallback"
    report_text = build_fallback_report(req, risk_data, policy_data, blockchain_data_for_report, None)
    external_risk_score = 0.0

    client = get_client()
    if client is not None:
        try:
            broadcast_log_sync({"type": "info", "message": "[3/6] Generating AI report with multi-agent pipeline..."})
            coordinator = create_root_coordinator(client)
            response = coordinator.generate_content(make_audit_prompt(req))
            report_text = response.text
            
            # Extract external_risk_score from coordinator context if available
            if hasattr(coordinator, 'context') and 'external_risk_score' in coordinator.context.state:
                external_risk_score = coordinator.context.state.get('external_risk_score', 0.0)
            
            report_source = "groq-llama"
            broadcast_log_sync({"type": "success", "message": "✓ AI report generated successfully"})
        except Exception as e:
            broadcast_log_sync({"type": "warning", "message": f"⚠ AI report failed, using fallback: {str(e)[:100]}"})
            report_source = "deterministic-fallback"
    else:
        broadcast_log_sync({"type": "warning", "message": "⚠ AI client unavailable, using deterministic report"})

    # Blockchain: Register report hash
    broadcast_log_sync({"type": "info", "message": "[4/6] Recording report hash on blockchain..."})
    try:
        report_hash_record = bc.register_report_hash(
            supplier_name=req.supplier_name,
            score_anchor_tx=score_anchor.get("tx_id") if score_anchor else None,
            hitl_decision_tx=None,
            report_text=report_text
        )
    except Exception as e:
        broadcast_log_sync({"type": "error", "message": f"✗ Report hash exception: {str(e)}"})
        report_hash_record = None
    
    if report_hash_record:
        broadcast_log_sync({"type": "success", "message": f"✓ Report hash: {report_hash_record['verification_code']}"})
        # Update blockchain data for report
        blockchain_data_for_report["report_hash"] = report_hash_record.get("report_hash")
        blockchain_data_for_report["verification_code"] = report_hash_record.get("verification_code")
        blockchain_data_for_report["report_tx"] = report_hash_record.get("tx_id") or report_hash_record.get("local_id")
    else:
        broadcast_log_sync({"type": "error", "message": "✗ Report hash failed - returned None"})
        report_hash_record = {"local_id": "FALLBACK", "on_chain": False}

    broadcast_log_sync({"type": "info", "message": "[5/6] Finalizing audit results..."})
    result = {
        "job_id": f"JOB-{uuid4().hex[:8].upper()}",
        "audit_id": f"AUD-{uuid4().hex[:10].upper()}",
        "timestamp": datetime.now(timezone.utc).isoformat(),
        "supplier_name": req.supplier_name,
        "emissions": req.emissions,
        "violations": req.violations,
        "notes": req.notes,
        "sector": risk_data.get("sector", "General Industry"),
        "sector_key": risk_data.get("sector_key", "default"),
        "production_volume": req.production_volume,
        "production_unit": req.production_unit,
        "registry_id": req.registry_id,
        "emissions_intensity": risk_data.get("emissions_intensity"),
        "prorata_progress": risk_data.get("prorata_progress", 0.0),
        "baseline_year": req.baseline_year,
        "target_year": req.target_year,
        "risk_score": risk_data["risk_score"],
        "classification": risk_data["classification"],
        "emissions_score": risk_data["emissions_score"],
        "violations_score": risk_data["violations_score"],
        "external_risk_score": external_risk_score,
        "policy_decision": policy_data["decision"],
        "human_approval_required": policy_data["human_approval_required"],
        "policy_reason": policy_data["reason"],
        "recommended_action": policy_data["recommended_action"],
        "report_text": report_text,
        "report_source": report_source,
        "download_links": {},
        "status": "pending_approval" if policy_data["human_approval_required"] else "completed",
        "blockchain": {
            "score_tx": score_anchor.get("tx_id") or score_anchor.get("local_id") if score_anchor else None,
            "score_hash": score_anchor.get("data_hash") if score_anchor else None,
            "report_tx": report_hash_record.get("tx_id") or report_hash_record.get("local_id") if report_hash_record else None,
            "report_hash": report_hash_record.get("report_hash") if report_hash_record else None,
            "verification_code": report_hash_record.get("verification_code") if report_hash_record else None,
            "on_chain": score_anchor.get("on_chain", False) if score_anchor else False
        }
    }

    # ── Carbon Credit Scoring (runs on every audit) ───────────────
    broadcast_log_sync({"type": "info", "message": "[6/6] Calculating carbon credits..."})
    credit_data_for_report = None
    try:
        credit_result = calculate_carbon_credits(result)
        result["carbon_credits"] = credit_result
        
        # Record credits on blockchain
        bc = get_blockchain_client()
        credit_record = bc.record_carbon_credits(
            supplier_name=req.supplier_name,
            audit_id=result["audit_id"],
            credits_earned=credit_result["credits_earned"],
            badges_earned=credit_result["badges_earned"],
            total_credits=credit_result["new_total"],
            esg_score=risk_data["risk_score"],
            streak_bonus=credit_result["streak_bonus"],
            improvement_bonus=credit_result["improvement_bonus"],
        )
        
        # Add credit transaction to blockchain data
        blockchain_data_for_report["credit_tx"] = credit_record.get("tx_id") or credit_record.get("local_id")
        credit_data_for_report = credit_result
        
        # Regenerate report with all blockchain and credit details
        report_text = build_fallback_report(req, risk_data, policy_data, blockchain_data_for_report, credit_data_for_report)
        
        broadcast_log_sync({
            "type": "success",
            "message": (
                f"✓ Credits: +{credit_result['total_credits']} "
                f"(total: {credit_result['new_total']}) | "
                f"Badges: {credit_result['badges_earned']}"
            ),
        })
    except Exception as e:
        broadcast_log_sync({"type": "warning", "message": f"⚠ Credit calculation failed: {str(e)[:80]}"})
        result["carbon_credits"] = None
        # Regenerate report without credits
        report_text = build_fallback_report(req, risk_data, policy_data, blockchain_data_for_report, None)

    # HITL Workflow Pause: If human approval required, save to pending queue
    if policy_data["human_approval_required"]:
        broadcast_log_sync({"type": "warning", "message": "🚨 CRITICAL RISK - Audit paused for human approval"})
        result["status"] = "pending_approval"
        result["approval_status"] = "pending"
        result["approver_name"] = None
        result["approval_notes"] = None
        result["approval_timestamp"] = None
        result["blockchain"]["hitl_tx"] = None
    else:
        broadcast_log_sync({"type": "success", "message": f"✓ Audit complete for {req.supplier_name}"})
    
    return result


def _write_pdf(pdf_path: Path, result: Dict[str, Any]) -> None:
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import inch
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, PageBreak
    from reportlab.lib import colors
    from reportlab.lib.enums import TA_LEFT, TA_CENTER
    
    doc = SimpleDocTemplate(str(pdf_path), pagesize=A4, topMargin=0.75*inch, bottomMargin=0.75*inch, leftMargin=0.75*inch, rightMargin=0.75*inch)
    styles = getSampleStyleSheet()
    story = []
    
    # Title
    title_style = ParagraphStyle(
        'CustomTitle',
        parent=styles['Heading1'],
        fontSize=18,
        textColor=colors.HexColor('#0d7c66'),
        spaceAfter=12,
        alignment=TA_CENTER
    )
    story.append(Paragraph(f"CfoE Audit Report", title_style))
    story.append(Spacer(1, 0.2*inch))
    
    # Summary Table with word wrapping
    summary_data = [
        ['Audit ID:', Paragraph(result['audit_id'], styles['Normal'])],
        ['Job ID:', Paragraph(result['job_id'], styles['Normal'])],
        ['Timestamp:', Paragraph(result['timestamp'][:19], styles['Normal'])],
        ['Supplier:', Paragraph(result['supplier_name'], styles['Normal'])],
        ['Emissions:', Paragraph(f"{result['emissions']} tons CO2", styles['Normal'])],
        ['Violations:', Paragraph(str(result['violations']), styles['Normal'])],
        ['Risk Score:', Paragraph(f"{result['risk_score']} ({result['classification']})", styles['Normal'])],
    ]
    
    summary_table = Table(summary_data, colWidths=[1.5*inch, 4.5*inch])
    summary_table.setStyle(TableStyle([
        ('BACKGROUND', (0, 0), (0, -1), colors.HexColor('#f0f0f0')),
        ('TEXTCOLOR', (0, 0), (-1, -1), colors.black),
        ('ALIGN', (0, 0), (0, -1), 'RIGHT'),
        ('ALIGN', (1, 0), (1, -1), 'LEFT'),
        ('FONTNAME', (0, 0), (0, -1), 'Helvetica-Bold'),
        ('FONTSIZE', (0, 0), (-1, -1), 9),
        ('GRID', (0, 0), (-1, -1), 0.5, colors.grey),
        ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
        ('TOPPADDING', (0, 0), (-1, -1), 8),
        ('BOTTOMPADDING', (0, 0), (-1, -1), 8),
        ('LEFTPADDING', (0, 0), (-1, -1), 8),
        ('RIGHTPADDING', (0, 0), (-1, -1), 8),
    ]))
    story.append(summary_table)
    story.append(Spacer(1, 0.3*inch))
    
    # Policy Section
    policy_style = ParagraphStyle('PolicyHeading', parent=styles['Heading2'], fontSize=14, textColor=colors.HexColor('#0d7c66'))
    story.append(Paragraph("Policy Enforcement", policy_style))
    story.append(Spacer(1, 0.1*inch))
    
    policy_data = [
        ['Decision:', Paragraph(result['policy_decision'], styles['Normal'])],
        ['Reason:', Paragraph(result['policy_reason'], styles['Normal'])],
        ['Recommended Action:', Paragraph(result['recommended_action'], styles['Normal'])],
    ]
    
    policy_table = Table(policy_data, colWidths=[1.5*inch, 4.5*inch])
    policy_table.setStyle(TableStyle([
        ('BACKGROUND', (0, 0), (0, -1), colors.HexColor('#f0f0f0')),
        ('TEXTCOLOR', (0, 0), (-1, -1), colors.black),
        ('ALIGN', (0, 0), (0, -1), 'RIGHT'),
        ('ALIGN', (1, 0), (1, -1), 'LEFT'),
        ('FONTNAME', (0, 0), (0, -1), 'Helvetica-Bold'),
        ('FONTSIZE', (0, 0), (-1, -1), 9),
        ('GRID', (0, 0), (-1, -1), 0.5, colors.grey),
        ('VALIGN', (0, 0), (-1, -1), 'TOP'),
        ('TOPPADDING', (0, 0), (-1, -1), 8),
        ('BOTTOMPADDING', (0, 0), (-1, -1), 8),
        ('LEFTPADDING', (0, 0), (-1, -1), 8),
        ('RIGHTPADDING', (0, 0), (-1, -1), 8),
    ]))
    story.append(policy_table)
    story.append(Spacer(1, 0.3*inch))
    
    # Executive Report
    story.append(Paragraph("Executive Report", policy_style))
    story.append(Spacer(1, 0.1*inch))
    
    # Format report text with proper line breaks and structure
    report_style = ParagraphStyle(
        'ReportBody', 
        parent=styles['BodyText'], 
        fontSize=9, 
        leading=13,
        alignment=TA_LEFT,
        leftIndent=0,
        rightIndent=0,
        spaceAfter=6
    )
    
    # Split report into lines and format properly
    report_lines = result['report_text'].split('\n')
    for line in report_lines:
        line = line.strip()
        if line:
            # Replace special characters that might cause issues
            line = line.replace('━', '-')
            line = line.replace('•', '*')
            
            # Check if it's a section header (numbered or all caps)
            if line and (line[0].isdigit() or line.isupper() or line.startswith('---')):
                header_style = ParagraphStyle(
                    'SectionHeader',
                    parent=styles['Heading3'],
                    fontSize=10,
                    textColor=colors.HexColor('#0d7c66'),
                    spaceAfter=4,
                    spaceBefore=8,
                    fontName='Helvetica-Bold'
                )
                story.append(Paragraph(line, header_style))
            else:
                story.append(Paragraph(line, report_style))
    
    doc.build(story)


def _write_docx(docx_path: Path, result: Dict[str, Any]) -> None:
    doc = Document()
    doc.add_heading(f"CfoE Audit Report - {result['audit_id']}", level=1)
    doc.add_paragraph(f"Job ID: {result['job_id']}")
    doc.add_paragraph(f"Timestamp: {result['timestamp']}")
    doc.add_paragraph(f"Supplier: {result['supplier_name']}")
    doc.add_paragraph(f"Emissions: {result['emissions']}")
    doc.add_paragraph(f"Violations: {result['violations']}")
    doc.add_paragraph(f"Risk Score: {result['risk_score']} ({result['classification']})")
    doc.add_paragraph(f"Policy Decision: {result['policy_decision']}")
    doc.add_paragraph(f"Policy Reason: {result['policy_reason']}")
    doc.add_paragraph(f"Recommended Action: {result['recommended_action']}")
    doc.add_heading("Executive Report", level=2)
    doc.add_paragraph(result["report_text"])
    doc.save(str(docx_path))


def export_audit_files(result: Dict[str, Any]) -> Dict[str, str]:
    ensure_storage()

    safe_stem = result["audit_id"].lower()
    job_dir = OUTPUT_DIR / result["job_id"].lower()
    job_dir.mkdir(parents=True, exist_ok=True)

    txt_path = job_dir / f"{safe_stem}.txt"
    docx_path = job_dir / f"{safe_stem}.docx"
    pdf_path = job_dir / f"{safe_stem}.pdf"

    # Save raw report text for verification
    txt_path.write_text(result["report_text"], encoding="utf-8")

    # Try to create PDF, but don't fail if it errors
    pdf_created = False
    try:
        _write_pdf(pdf_path, result)
        pdf_created = True
    except Exception as e:
        print(f"[WARNING] PDF generation failed: {e}")
        print("[INFO] Continuing without PDF export")

    # Try to create DOCX
    docx_created = False
    try:
        _write_docx(docx_path, result)
        docx_created = True
    except Exception as e:
        print(f"[WARNING] DOCX generation failed: {e}")
        print("[INFO] Continuing without DOCX export")

    with lock:
        with OUTPUT_CSV_PATH.open("a", encoding="utf-8", newline="") as csv_file:
            writer = csv.DictWriter(
                csv_file,
                fieldnames=[
                    "audit_id",
                    "timestamp",
                    "supplier_name",
                    "emissions",
                    "violations",
                    "risk_score",
                    "classification",
                    "policy_decision",
                    "human_approval_required",
                    "report_source",
                ],
            )
            writer.writerow(
                {
                    "audit_id": result["audit_id"],
                    "timestamp": result["timestamp"],
                    "supplier_name": result["supplier_name"],
                    "emissions": result["emissions"],
                    "violations": result["violations"],
                    "risk_score": result["risk_score"],
                    "classification": result["classification"],
                    "policy_decision": result["policy_decision"],
                    "human_approval_required": result["human_approval_required"],
                    "report_source": result["report_source"],
                }
            )

    links = {"txt": f"/outputs/{result['job_id'].lower()}/{txt_path.name}"}
    if pdf_created:
        links["pdf"] = f"/api/audits/{result['audit_id']}/pdf/download"
    if docx_created:
        links["docx"] = f"/outputs/{result['job_id'].lower()}/{docx_path.name}"
    
    # Always include TXT in download dialog
    return links


# ── Wallet state (must be defined before startup_event uses it) ───

_wallet_state: Dict[str, Any] = {"connected": False, "address": None}


@app.on_event("startup")
def startup_event() -> None:
    ensure_storage()

    # Sync _wallet_state with .env auto-connected wallet
    # This ensures the frontend sees the correct state on fresh device setups
    bc = get_blockchain_client()
    if bc.wallet_connected and bc.address:
        _wallet_state["connected"] = True
        _wallet_state["address"] = bc.address
        print(f"  [Startup] Auto-synced wallet state from .env: {bc.address[:16]}...")

    # ── Part 1: Initialize agent wallets ──────────────────────────────────
    try:
        from agents.agent_wallets import initialize_agent_wallets
        wallet_info = initialize_agent_wallets()
        for agent, info in wallet_info.items():
            addr = info.get("address", "N/A")
            bal = info.get("balance_algo", 0.0)
            print(f"  [AgentWallet] {agent}: {addr[:20]}... | {bal:.6f} ALGO")
    except Exception as _wex:
        print(f"  [AgentWallet] WARNING: Could not initialize agent wallets: {_wex}")


@app.get("/")
def serve_index() -> FileResponse:
    return FileResponse(WEB_DIR / "index.html")


@app.get("/simulator")
def serve_simulator() -> FileResponse:
    """Serve simulator dashboard on same port"""
    return FileResponse(SIMULATOR_DIR / "dashboard.html")


@app.post("/api/audit", response_model=AuditResponse)
async def create_audit(
    payload: AuditRequest,
    request: Request,
    x_payment: Optional[str] = Header(None, alias="X-Payment"),
    x_internal_secret: Optional[str] = Header(None, alias="X-Internal-Secret"),
) -> Dict[str, Any]:
    """POST /api/audit — X402 payment gate (Part 3).

    External callers must include a valid X-Payment header (0.05 ALGO).
    Internal simulator calls bypass the gate using the X-Internal-Secret header.
    """
    # Phase 2: Validate registry ID if provided
    if payload.registry_id and payload.registry_id.strip():
        validation = validate_registry_id(payload.registry_id)
        if not validation["valid"]:
            raise HTTPException(status_code=400, detail=validation["error"])

    # ── X402 Payment Gate (Part 3) ─────────────────────────────────────────
    is_internal = x_internal_secret == X402_INTERNAL_SECRET
    payment_info: Optional[Dict[str, Any]] = None

    if not is_internal:
        if not x_payment:
            # Return 402 with payment instructions
            bc = get_blockchain_client()
            auditor_addr = bc.address or ""
            try:
                from agents.x402_payments import build_payment_required_body
                body = build_payment_required_body(
                    receiver_address=auditor_addr,
                    amount_algo=AUDIT_PAYMENT_ALGO,
                    description=f"CfoE ESG Audit — {AUDIT_PAYMENT_ALGO} ALGO per audit",
                )
            except Exception:
                body = {
                    "x402Version": 1,
                    "error": "Payment required",
                    "payTo": auditor_addr,
                    "amount_algo": AUDIT_PAYMENT_ALGO,
                }
            return JSONResponse(status_code=402, content=body)

        # Validate the provided payment header
        try:
            from agents.x402_payments import validate_audit_payment
            bc = get_blockchain_client()
            auditor_addr = bc.address or ""
            valid, err_msg, payment_info = validate_audit_payment(
                x_payment_header=x_payment,
                auditor_address=auditor_addr,
                required_amount_algo=AUDIT_PAYMENT_ALGO,
            )
        except Exception as pex:
            valid, err_msg, payment_info = False, str(pex), None

        if not valid:
            raise HTTPException(status_code=402, detail=f"Invalid payment: {err_msg}")

    try:
        # Clear log queue before starting
        while not log_queue.empty():
            log_queue.get()

        result = run_audit(payload)
        result["download_links"] = export_audit_files(result)
    except Exception as exc:
        import traceback
        error_trace = traceback.format_exc()
        print(f"\n[ERROR] Audit failed with exception:")
        print(error_trace)
        raise HTTPException(status_code=500, detail=f"Audit failed: {str(exc)}") from exc

    # Attach payment info to history record (Part 3)
    if payment_info:
        result["x402_payment"] = payment_info
        try:
            from agents.x402_payments import record_payment
            record_payment(
                agent_name="auditor",
                amount_algo=AUDIT_PAYMENT_ALGO,
                service="audit",
                tx_id=payment_info.get("tx_id"),
                direction="incoming",
                status="confirmed",
                audit_id=result.get("audit_id"),
            )
        except Exception:
            pass

    # If human approval required, save to pending queue instead of history
    if result.get("human_approval_required", False):
        pending = load_pending()
        pending.insert(0, result)
        save_pending(pending)
    else:
        # Auto-approved audits go directly to history
        history = load_history()
        history.insert(0, result)
        save_history(history[:500])

    return result


@app.get("/api/audits")
def list_audits(limit: int = 100) -> Dict[str, Any]:
    history = load_history()
    return {"items": history[:max(1, min(limit, 500))], "count": len(history)}


@app.get("/api/approvals")
def list_pending_approvals() -> Dict[str, Any]:
    """Get all audits pending human approval"""
    pending = load_pending()
    return {"items": pending, "count": len(pending)}


@app.post("/api/approvals/{audit_id}/approve")
def approve_audit(audit_id: str, approval: ApprovalRequest) -> Dict[str, Any]:
    """Approve a pending audit and move it to history"""
    if approval.audit_id != audit_id:
        raise HTTPException(status_code=400, detail="Audit ID mismatch")
    
    if approval.decision != "approve":
        raise HTTPException(status_code=400, detail="Use /reject endpoint for rejections")
    
    pending = load_pending()
    audit = next((x for x in pending if x.get("audit_id") == audit_id), None)
    
    if audit is None:
        raise HTTPException(status_code=404, detail="Pending audit not found")
    
    # Blockchain: Record HITL decision
    bc = get_blockchain_client()
    hitl_record = bc.record_hitl_decision(
        supplier_name=audit["supplier_name"],
        score_anchor_tx=audit.get("blockchain", {}).get("score_tx"),
        approved=True,
        risk_score=audit["risk_score"],
        decision=audit["policy_decision"],
        reason=audit["policy_reason"],
        recommended_action=audit["recommended_action"]
    )
    
    # Update audit with approval info
    audit["status"] = "completed"
    audit["approval_status"] = "approved"
    audit["approver_name"] = approval.approver_name
    audit["approval_notes"] = approval.approval_notes
    audit["approval_timestamp"] = datetime.now(timezone.utc).isoformat()
    if "blockchain" not in audit:
        audit["blockchain"] = {}
    audit["blockchain"]["hitl_tx"] = hitl_record.get("tx_id") or hitl_record.get("local_id")
    
    # Remove from pending and add to history
    pending = [x for x in pending if x.get("audit_id") != audit_id]
    save_pending(pending)
    
    history = load_history()
    history.insert(0, audit)
    save_history(history[:500])
    
    return {"status": "approved", "audit": audit}


@app.post("/api/approvals/{audit_id}/reject")
def reject_audit(audit_id: str, approval: ApprovalRequest) -> Dict[str, Any]:
    """Reject a pending audit"""
    if approval.audit_id != audit_id:
        raise HTTPException(status_code=400, detail="Audit ID mismatch")
    
    if approval.decision != "reject":
        raise HTTPException(status_code=400, detail="Use /approve endpoint for approvals")
    
    pending = load_pending()
    audit = next((x for x in pending if x.get("audit_id") == audit_id), None)
    
    if audit is None:
        raise HTTPException(status_code=404, detail="Pending audit not found")
    
    # Blockchain: Record HITL decision
    bc = get_blockchain_client()
    hitl_record = bc.record_hitl_decision(
        supplier_name=audit["supplier_name"],
        score_anchor_tx=audit.get("blockchain", {}).get("score_tx"),
        approved=False,
        risk_score=audit["risk_score"],
        decision=audit["policy_decision"],
        reason=audit["policy_reason"],
        recommended_action=audit["recommended_action"]
    )
    
    # Update audit with rejection info
    audit["status"] = "rejected"
    audit["approval_status"] = "rejected"
    audit["approver_name"] = approval.approver_name
    audit["approval_notes"] = approval.approval_notes
    audit["approval_timestamp"] = datetime.now(timezone.utc).isoformat()
    if "blockchain" not in audit:
        audit["blockchain"] = {}
    audit["blockchain"]["hitl_tx"] = hitl_record.get("tx_id") or hitl_record.get("local_id")
    
    # Remove from pending and add to history (with rejected status)
    pending = [x for x in pending if x.get("audit_id") != audit_id]
    save_pending(pending)
    
    history = load_history()
    history.insert(0, audit)
    save_history(history[:500])
    
    return {"status": "rejected", "audit": audit}


@app.delete("/api/audits")
def clear_audits() -> Dict[str, Any]:
    save_history([])
    return {"status": "ok"}


@app.delete("/api/approvals")
def clear_pending_approvals() -> Dict[str, Any]:
    """Clear all pending approvals"""
    save_pending([])
    return {"status": "ok"}


@app.get("/api/metrics")
def metrics() -> Dict[str, Any]:
    history = load_history()
    if not history:
        return {
            "total_audits": 0,
            "avg_risk_score": 0,
            "critical_rate": 0,
            "classifications": {"Low Risk": 0, "Moderate Risk": 0, "Critical Risk": 0},
        }

    total = len(history)
    avg_score = sum(item["risk_score"] for item in history) / total
    critical = sum(1 for item in history if item["classification"] == "Critical Risk")
    counts = {"Low Risk": 0, "Moderate Risk": 0, "Critical Risk": 0}
    for item in history:
        counts[item["classification"]] = counts.get(item["classification"], 0) + 1

    return {
        "total_audits": total,
        "avg_risk_score": round(avg_score, 3),
        "critical_rate": round((critical / total) * 100, 1),
        "classifications": counts,
    }


@app.get("/api/audits/{audit_id}/pdf/view")
def view_pdf(audit_id: str) -> FileResponse:
    history = load_history()
    item = next((x for x in history if x.get("audit_id") == audit_id), None)
    if item is None:
        raise HTTPException(status_code=404, detail="Audit not found")

    job_id = item.get("job_id", "")
    if not job_id:
        raise HTTPException(status_code=404, detail="Job export not found")

    pdf_name = f"{audit_id.lower()}.pdf"
    pdf_path = OUTPUT_DIR / job_id.lower() / pdf_name
    if not pdf_path.exists():
        raise HTTPException(status_code=404, detail="PDF file not found")

    return FileResponse(
        str(pdf_path),
        media_type="application/pdf",
        headers={"Content-Disposition": f"inline; filename={pdf_name}"},
    )


@app.get("/api/audits/{audit_id}/pdf/download")
def download_pdf(audit_id: str) -> FileResponse:
    history = load_history()
    item = next((x for x in history if x.get("audit_id") == audit_id), None)
    if item is None:
        raise HTTPException(status_code=404, detail="Audit not found")

    job_id = item.get("job_id", "")
    if not job_id:
        raise HTTPException(status_code=404, detail="Job export not found")

    pdf_name = f"{audit_id.lower()}.pdf"
    pdf_path = OUTPUT_DIR / job_id.lower() / pdf_name
    if not pdf_path.exists():
        raise HTTPException(status_code=404, detail="PDF file not found")

    return FileResponse(
        str(pdf_path),
        media_type="application/pdf",
        filename=pdf_name,
        headers={"Content-Disposition": f"attachment; filename={pdf_name}"},
    )


@app.get("/api/blockchain/status")
def blockchain_status() -> Dict[str, Any]:
    """Get blockchain connection status and statistics"""
    bc = get_blockchain_client()
    tm = get_token_manager()
    
    # Force refresh connection to ensure we have latest wallet info
    if not bc.connected:
        bc.connect()
    
    # Re-check wallet connection from .env if not already connected
    if not bc.wallet_connected:
        import os
        try:
            from algosdk import account
            env_key = os.getenv("ALGORAND_PRIVATE_KEY")
            if env_key:
                try:
                    bc.address = account.address_from_private_key(env_key)
                    bc.private_key = env_key
                    bc.wallet_connected = True
                except Exception:
                    pass
        except ImportError:
            pass
    
    # Keep _wallet_state in sync with actual blockchain client state
    # This is critical for fresh device setups where .env auto-connects
    # the blockchain client but _wallet_state was never updated
    if bc.wallet_connected and bc.address:
        _wallet_state["connected"] = True
        _wallet_state["address"] = bc.address
    
    balance_info = bc.get_balance()
    history = bc.get_audit_history()
    
    # Show full address if wallet connected, otherwise show N/A
    display_address = bc.address if bc.wallet_connected else "N/A"
    
    # Get token information
    total_issued = sum(r["carbon_credits"] for r in tm.issued_credits)
    total_retired = sum(r["carbon_credits"] for r in tm.retired_credits)
    
    # Get token balance (returns dict with tokens and carbon_credits)
    token_balance_info = tm.get_credit_balance(bc.address) if bc.address else {"tokens": 0.0, "carbon_credits": 0.0}
    
    # Determine overall connection status
    is_connected = bc.wallet_connected  # Show connected if wallet is connected
    network_status = "Algorand Testnet" if bc.connected else "Offline"
    
    return {
        "connected": is_connected,
        "address": display_address,
        "balance": balance_info.get("balance_algo", 0),
        "network": network_status,
        "wallet": _wallet_state,
        "wallet_connected": bc.wallet_connected,
        "token_id": tm.carbon_credit_asset_id,
        "token_balance": token_balance_info.get("tokens", 0),
        "token_supply": total_issued - total_retired,
        "credits_issued": total_issued,
        "credits_retired": total_retired,
        "score_anchors": len(history.get("score_anchors", [])),
        "hitl_decisions": len(history.get("hitl_decisions", [])),
        "report_hashes": len(history.get("report_hashes", [])),
        "total_blockchain_records": (
            len(history.get("score_anchors", []))
            + len(history.get("hitl_decisions", []))
            + len(history.get("report_hashes", []))
        ),
    }



@app.get("/api/registry/validate/{registry_id}")
def validate_registry(registry_id: str) -> Dict[str, Any]:
    """Validate entity registry ID"""
    return validate_registry_id(registry_id)


@app.get("/api/registry/entity/{registry_id}")
def get_entity(registry_id: str) -> Dict[str, Any]:
    """Get entity information by registry ID"""
    entity = get_entity_info(registry_id)
    if entity is None:
        raise HTTPException(status_code=404, detail="Entity not found")
    return entity


@app.get("/api/trajectory/{supplier_name}")
def get_trajectory(supplier_name: str) -> Dict[str, Any]:
    """Get multi-year compliance trajectory for a supplier"""
    history = load_history()
    return calculate_trajectory(supplier_name, history)


@app.get("/api/trajectory/{supplier_name}/compliance")
def get_compliance_trajectory(supplier_name: str, baseline_year: int = 2023, target_year: int = 2027) -> Dict[str, Any]:
    """Check if supplier is on track to meet compliance goals"""
    history = load_history()
    return check_compliance_trajectory(supplier_name, history, baseline_year, target_year)


# ── Carbon Credit endpoints ──────────────────────────────────────

@app.get("/api/credits/{supplier_id}")
def get_credits(supplier_id: str) -> Dict[str, Any]:
    """Return the full carbon credit history for a supplier."""
    data = get_supplier_credits(supplier_id)
    if data is None:
        raise HTTPException(status_code=404, detail="Supplier not found in credit ledger")
    return data


@app.get("/api/credits/{supplier_id}/history")
def get_credit_history(supplier_id: str) -> Dict[str, Any]:
    """Return rich credit history for charts and timelines:
    credits per audit (sparkline), badge timeline, streak history,
    and ESG score trend."""
    data = get_supplier_credit_history(supplier_id)
    if data is None:
        raise HTTPException(status_code=404, detail="Supplier not found in credit ledger")
    return data


@app.get("/api/leaderboard")
def leaderboard() -> Dict[str, Any]:
    """Return all suppliers sorted by total credits descending,
    with their badges and latest ESG score."""
    board = get_leaderboard()
    return {"items": board, "count": len(board)}


# ── Wallet endpoints ──────────────────────────────────────────────

class WalletConnectRequest(BaseModel):
    address: str = Field(min_length=1, max_length=128)


@app.get("/api/wallet/status")
def wallet_status() -> Dict[str, Any]:
    """Return the current wallet connection state."""
    # Sync with actual blockchain client state (handles .env auto-connect on fresh devices)
    bc = get_blockchain_client()
    if bc.wallet_connected and bc.address:
        _wallet_state["connected"] = True
        _wallet_state["address"] = bc.address
    return _wallet_state


@app.post("/api/wallet/connect")
def wallet_connect(payload: WalletConnectRequest) -> Dict[str, Any]:
    """Register a wallet address from the frontend."""
    _wallet_state["connected"] = True
    _wallet_state["address"] = payload.address

    # Connect wallet to blockchain client
    bc = get_blockchain_client()
    bc.set_wallet_address(payload.address)

    # Part 2: Flush any pending CCC mints for this wallet.
    # We match on address since we may not know the supplier_id at connect time.
    # Try to resolve supplier_id from audit history.
    flushed = []
    try:
        from onchain_ops import flush_pending_mints
        history = load_history()
        matched_ids = {
            rec["supplier_name"].strip().lower().replace(" ", "_")
            for rec in history
            if rec.get("supplier_name")
        }
        for sid in matched_ids:
            processed = flush_pending_mints(payload.address, sid)
            flushed.extend(processed)
    except Exception as _fe:
        pass

    return {"status": "ok", "address": payload.address, "pending_mints_flushed": len(flushed)}


@app.post("/api/wallet/disconnect")
def wallet_disconnect() -> Dict[str, Any]:
    """Clear the wallet connection."""
    _wallet_state["connected"] = False
    _wallet_state["address"] = None
    
    # Disconnect wallet from blockchain client
    bc = get_blockchain_client()
    bc.disconnect_wallet()
    
    return {"status": "ok"}


# ── Carbon Credit Token endpoints ─────────────────────────────────

class TokenCreateRequest(BaseModel):
    total_credits: int = Field(default=10_000_000, ge=1000)
    unit_name: str = Field(default="CCT", max_length=8)
    asset_name: str = Field(default="CfoE Carbon Credit", max_length=32)


class CreditIssueRequest(BaseModel):
    recipient_address: str = Field(min_length=58, max_length=58)
    amount: float = Field(gt=0)
    reason: str = Field(max_length=200)
    audit_id: Optional[str] = None


class TokenOptInRequest(BaseModel):
    asset_id: int = Field(gt=0)


class CreditRetireRequest(BaseModel):
    amount: float = Field(gt=0)
    reason: str = Field(max_length=200)
    beneficiary: str = Field(max_length=100)


class CreditTransferRequest(BaseModel):
    recipient_address: str = Field(min_length=58, max_length=58)
    amount: float = Field(gt=0)
    reason: str = Field(max_length=200)
    audit_id: Optional[str] = None


class NFTCreateRequest(BaseModel):
    supplier_name: str
    audit_id: str
    risk_score: float
    classification: str
    emissions: float
    metadata_url: str = ""


class SetAssetIdRequest(BaseModel):
    asset_id: int


@app.post("/api/tokens/set-asset-id")
def set_asset_id(payload: SetAssetIdRequest) -> Dict[str, Any]:
    """Save/update the asset ID in the token_state.json file."""
    from carbon_token_manager import _save_asset_id
    try:
        _save_asset_id(payload.asset_id)
        
        tm = get_token_manager()
        tm.carbon_credit_asset_id = payload.asset_id
        
        return {
            "status": "success",
            "asset_id": payload.asset_id,
            "message": f"Asset ID {payload.asset_id} saved successfully"
        }
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Failed to save asset ID: {str(e)}")


@app.post("/api/tokens/optin")
def optin_to_token(payload: TokenOptInRequest) -> Dict[str, Any]:
    """Opt-in the server wallet to receive carbon credit tokens."""
    tm = get_token_manager()
    bc = get_blockchain_client()
    
    if not bc.connected or not bc.wallet_connected:
        raise HTTPException(status_code=400, detail="Wallet not connected")
    
    try:
        tx_id = tm.optin_to_asset(payload.asset_id)
        
        if tx_id:
            return {
                "status": "success",
                "tx_id": tx_id,
                "message": f"Successfully opted-in to asset {payload.asset_id}"
            }
        else:
            raise HTTPException(status_code=500, detail="Opt-in transaction failed")
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Opt-in failed: {str(e)}")


@app.post("/api/tokens/create")
def create_carbon_token(payload: TokenCreateRequest) -> Dict[str, Any]:
    """Create a new carbon credit token (ASA)."""
    tm = get_token_manager()
    asset_id = tm.create_carbon_credit_token(
        total_credits=payload.total_credits,
        unit_name=payload.unit_name,
        asset_name=payload.asset_name,
    )
    
    if asset_id:
        return {
            "status": "success",
            "asset_id": asset_id,
            "message": f"Carbon credit token created: {asset_id}"
        }
    else:
        raise HTTPException(status_code=500, detail="Token creation failed")


@app.post("/api/tokens/issue")
def issue_carbon_credits(payload: CreditIssueRequest) -> Dict[str, Any]:
    """Issue carbon credits to a recipient."""
    tm = get_token_manager()
    
    if not tm.carbon_credit_asset_id:
        raise HTTPException(
            status_code=400,
            detail="No carbon credit token exists yet. Create one first via POST /api/tokens/create"
        )
    
    tx_id = tm.issue_credits(
        recipient_address=payload.recipient_address,
        carbon_credits=payload.amount,
        reason=payload.reason,
        audit_id=payload.audit_id,
    )
    
    if tx_id:
        return {
            "status": "success",
            "tx_id": tx_id,
            "amount": payload.amount,
            "message": f"Issued {payload.amount} tons CO2eq"
        }
    else:
        raise HTTPException(status_code=500, detail="Credit issuance failed")


@app.post("/api/tokens/retire")
def retire_carbon_credits(payload: CreditRetireRequest) -> Dict[str, Any]:
    """Retire (burn) carbon credits permanently."""
    tm = get_token_manager()
    tx_id = tm.retire_credits(
        carbon_credits=payload.amount,
        reason=payload.reason,
        beneficiary=payload.beneficiary,
    )
    
    if tx_id:
        return {
            "status": "success",
            "tx_id": tx_id,
            "amount": payload.amount,
            "message": f"Retired {payload.amount} tons CO2eq permanently"
        }
    else:
        raise HTTPException(status_code=500, detail="Credit retirement failed")


@app.post("/api/tokens/transfer")
def transfer_carbon_credits(payload: CreditTransferRequest) -> Dict[str, Any]:
    """Transfer carbon credits to another address."""
    tm = get_token_manager()

    if not tm.carbon_credit_asset_id:
        raise HTTPException(
            status_code=400,
            detail="No carbon credit token exists yet. Create one first via POST /api/tokens/create"
        )

    tx_id = tm.transfer_credits(
        recipient_address=payload.recipient_address,
        carbon_credits=payload.amount,
        reason=payload.reason,
        audit_id=payload.audit_id,
    )
    
    if tx_id:
        return {
            "status": "success",
            "tx_id": tx_id,
            "amount": payload.amount,
            "recipient": payload.recipient_address,
            "message": f"Transferred {payload.amount} tons CO2eq"
        }
    else:
        raise HTTPException(status_code=500, detail="Credit transfer failed")


@app.post("/api/tokens/nft/create")
def create_audit_nft(payload: NFTCreateRequest) -> Dict[str, Any]:
    """Create an audit certificate NFT."""
    tm = get_token_manager()
    asset_id = tm.create_audit_certificate_nft(
        supplier_name=payload.supplier_name,
        audit_id=payload.audit_id,
        risk_score=payload.risk_score,
        classification=payload.classification,
        emissions=payload.emissions,
        metadata_url=payload.metadata_url,
    )
    
    if asset_id:
        return {
            "status": "success",
            "asset_id": asset_id,
            "message": f"Audit certificate NFT created: {asset_id}"
        }
    else:
        raise HTTPException(status_code=500, detail="NFT creation failed")


@app.get("/api/tokens/balance/{address}")
def get_token_balance(address: str) -> Dict[str, Any]:
    """Get carbon credit balance for an address."""
    tm = get_token_manager()
    tm.refresh_ledgers()
    balance = tm.get_credit_balance(address)
    
    return {
        "address": address,
        "balance": balance,
        "asset_id": tm.carbon_credit_asset_id,
    }


@app.get("/api/tokens/summary")
def get_token_summary() -> Dict[str, Any]:
    """Get summary of all token operations."""
    tm = get_token_manager()
    tm.refresh_ledgers()
    
    total_issued = sum(r.get("carbon_credits", 0) for r in tm.issued_credits)
    total_retired = sum(r.get("carbon_credits", 0) for r in tm.retired_credits)
    
    return {
        "asset_id": tm.carbon_credit_asset_id,
        "total_issued": total_issued,
        "total_retired": total_retired,
        "circulating_supply": total_issued - total_retired,
        "issuance_count": len(tm.issued_credits),
        "retirement_count": len(tm.retired_credits),
        "nft_count": len(tm.audit_nfts),
        "issued_credits": tm.issued_credits,
        "retired_credits": tm.retired_credits,
        "audit_nfts": tm.audit_nfts,
    }


# ── Marketplace endpoints (Part 4) ─────────────────────────────────────────

class MarketplaceListRequest(BaseModel):
    supplier_id: str = Field(min_length=1, max_length=120)
    supplier_address: str = Field(min_length=58, max_length=58)
    amount_ccc: int = Field(gt=0)
    price_per_unit_micro_algo: int = Field(gt=0,
        description="Price per CCC token in micro-ALGO (1 ALGO = 1_000_000 micro-ALGO)")


class MarketplaceBuyRequest(BaseModel):
    listing_id: int = Field(gt=0)
    buyer_address: str = Field(min_length=58, max_length=58)
    buyer_supplier_id: str = Field(min_length=1, max_length=120)


@app.post("/api/marketplace/list")
def marketplace_list(payload: MarketplaceListRequest) -> Dict[str, Any]:
    """
    Supplier lists CCC tokens for sale.
    POST /api/marketplace/list
    Body: {supplier_id, supplier_address, amount_ccc, price_per_unit_micro_algo}
    """
    try:
        from onchain_ops import create_listing_offchain
        listing = create_listing_offchain(
            supplier_id=payload.supplier_id,
            supplier_address=payload.supplier_address,
            amount_ccc=payload.amount_ccc,
            price_per_unit_micro_algo=payload.price_per_unit_micro_algo,
        )
        return {
            "status": "listed",
            "listing": listing,
            "total_algo": (payload.amount_ccc * payload.price_per_unit_micro_algo) / 1_000_000,
        }
    except Exception as exc:
        raise HTTPException(status_code=500, detail=f"Listing failed: {str(exc)}")


@app.post("/api/marketplace/buy")
def marketplace_buy(payload: MarketplaceBuyRequest) -> Dict[str, Any]:
    """
    Buyer purchases a listed credit bundle.
    POST /api/marketplace/buy
    Body: {listing_id, buyer_address, buyer_supplier_id}

    Atomicity guarantee: payment and token transfer happen off-chain via
    algokit-utils AtomicTransactionComposer; this endpoint records the sale.
    """
    try:
        from onchain_ops import execute_buy_offchain
        result = execute_buy_offchain(
            listing_id=payload.listing_id,
            buyer_address=payload.buyer_address,
            buyer_supplier_id=payload.buyer_supplier_id,
        )
        if result.get("status") == "error":
            raise HTTPException(status_code=400, detail=result["reason"])
        return result
    except HTTPException:
        raise
    except Exception as exc:
        raise HTTPException(status_code=500, detail=f"Purchase failed: {str(exc)}")


@app.get("/api/marketplace/listings")
def marketplace_listings(active_only: bool = True) -> Dict[str, Any]:
    """Return all marketplace listings."""
    try:
        from onchain_ops import _load_listings
        store = _load_listings()
        listings = list(store["listings"].values())
        if active_only:
            listings = [l for l in listings if l.get("status") == "active"]
        return {"items": listings, "count": len(listings)}
    except Exception as exc:
        raise HTTPException(status_code=500, detail=str(exc))


# ── Staking endpoints (Part 5) ───────────────────────────────────────────────

class StakeRequest(BaseModel):
    supplier_id: str = Field(min_length=1, max_length=120)
    supplier_address: str = Field(min_length=58, max_length=58)
    amount_ccc: int = Field(gt=0)


class UnstakeRequest(BaseModel):
    supplier_id: str = Field(min_length=1, max_length=120)


@app.get("/api/staking/{supplier_id}")
def get_staking_status(supplier_id: str) -> Dict[str, Any]:
    """
    GET /api/staking/{supplier_id}
    Returns current stake position and pending yield for a supplier.
    """
    try:
        from onchain_ops import get_stake_status
        return get_stake_status(supplier_id)
    except Exception as exc:
        raise HTTPException(status_code=500, detail=str(exc))


@app.post("/api/staking/stake")
def staking_stake(payload: StakeRequest) -> Dict[str, Any]:
    """
    POST /api/staking/stake
    Lock CCC tokens for 30 days to earn 10% ALGO yield.
    Body: {supplier_id, supplier_address, amount_ccc}
    """
    try:
        from onchain_ops import stake_ccc
        result = stake_ccc(
            supplier_id=payload.supplier_id,
            supplier_address=payload.supplier_address,
            amount_ccc=payload.amount_ccc,
        )
        if result.get("status") == "error":
            raise HTTPException(status_code=400, detail=result["reason"])
        return result
    except HTTPException:
        raise
    except Exception as exc:
        raise HTTPException(status_code=500, detail=f"Stake failed: {str(exc)}")


@app.post("/api/staking/unstake")
def staking_unstake(payload: UnstakeRequest) -> Dict[str, Any]:
    """
    POST /api/staking/unstake
    Retrieve CCC tokens after 30-day lock period.
    Body: {supplier_id}
    """
    try:
        from onchain_ops import unstake_ccc
        result = unstake_ccc(payload.supplier_id)
        if result.get("status") == "error":
            raise HTTPException(status_code=400, detail=result["reason"])
        return result
    except HTTPException:
        raise
    except Exception as exc:
        raise HTTPException(status_code=500, detail=f"Unstake failed: {str(exc)}")


@app.post("/api/staking/claim-yield")
def staking_claim_yield(payload: UnstakeRequest) -> Dict[str, Any]:
    """
    POST /api/staking/claim-yield
    Claim 10% ALGO yield after lock period.
    Body: {supplier_id}
    """
    try:
        from onchain_ops import claim_yield_offchain
        result = claim_yield_offchain(payload.supplier_id)
        if result.get("status") == "error":
            raise HTTPException(status_code=400, detail=result["reason"])
        return result
    except HTTPException:
        raise
    except Exception as exc:
        raise HTTPException(status_code=500, detail=f"Claim yield failed: {str(exc)}")


# ── Compliance bonds query (Part 3) ──────────────────────────────────────────

@app.get("/api/bonds/{supplier_id}")
def get_bond_status(supplier_id: str) -> Dict[str, Any]:
    """Return current compliance bond status for a supplier."""
    try:
        from onchain_ops import _load_bonds
        bonds = _load_bonds()
        bond = bonds.get(supplier_id)
        if not bond:
            return {"supplier_id": supplier_id, "status": "no_bond"}
        return bond
    except Exception as exc:
        raise HTTPException(status_code=500, detail=str(exc))


@app.get("/api/bonds")
def list_bonds(active_only: bool = False) -> Dict[str, Any]:
    """Return all compliance bonds."""
    try:
        from onchain_ops import _load_bonds
        bonds = _load_bonds()
        items = list(bonds.values())
        if active_only:
            items = [b for b in items if b.get("status") == "active"]
        return {"items": items, "count": len(items)}
    except Exception as exc:
        raise HTTPException(status_code=500, detail=str(exc))



# ══════════════════════════════════════════════════════════════════════════
# Part 4 — ReportingAgent Pay-Per-Report Endpoints
# ══════════════════════════════════════════════════════════════════════════

@app.get("/api/report/{audit_id}")
async def get_encrypted_report(audit_id: str) -> Dict[str, Any]:
    """
    GET /api/report/{audit_id}

    Returns the encrypted report blob + payment instructions.
    The caller must first pay 0.02 ALGO to the ReportingAgent wallet
    and then call POST /api/report/{audit_id}/pay with the TX ID.
    """
    from agents.reporting_agent import get_report_blob  # type: ignore

    entry = get_report_blob(audit_id)
    if not entry:
        raise HTTPException(status_code=404, detail="Report not found for this audit_id")

    if entry.get("paid"):
        # Already paid — return decrypted report
        from agents.reporting_agent import get_decrypted_report  # type: ignore
        available, text, err = get_decrypted_report(audit_id)
        if available:
            return {"audit_id": audit_id, "paid": True, "report_text": text}
        raise HTTPException(status_code=500, detail=f"Decryption failed: {err}")

    # Not paid yet — return blob + payment instructions
    try:
        from agents.agent_wallets import get_agent_address  # type: ignore
        reporting_wallet = get_agent_address("reporting_agent")
    except Exception:
        reporting_wallet = None

    bc = get_blockchain_client()
    try:
        from agents.x402_payments import build_payment_required_body  # type: ignore
        instructions = build_payment_required_body(
            receiver_address=reporting_wallet or bc.address or "",
            amount_algo=REPORT_PAYMENT_ALGO,
            description="CfoE Audit Report — 0.02 ALGO for decrypted access",
        )
    except Exception:
        instructions = {
            "payTo": reporting_wallet or bc.address or "",
            "amount_algo": REPORT_PAYMENT_ALGO,
        }

    return {
        "audit_id": audit_id,
        "paid": False,
        "supplier_name": entry.get("supplier_name"),
        "encrypted_blob": entry.get("encrypted_blob"),
        "payment_instructions": instructions,
        "pay_url": f"/api/report/{audit_id}/pay",
        "message": f"Send {REPORT_PAYMENT_ALGO} ALGO to the ReportingAgent wallet, "
                   f"then POST to /api/report/{audit_id}/pay with the TX ID.",
    }


class ReportPayRequest(BaseModel):
    tx_id: str = Field(min_length=50, max_length=70)


@app.post("/api/report/{audit_id}/pay")
async def confirm_report_payment(audit_id: str, body: ReportPayRequest) -> Dict[str, Any]:
    """
    POST /api/report/{audit_id}/pay

    Body: { "tx_id": "<algorand_tx_id>" }

    Verifies the payment on-chain (via algod — never trusts header alone).
    On success, marks the report as paid and returns the decrypted content.
    Times out after 10 seconds and returns status "pending" if unconfirmed.
    """
    from agents.reporting_agent import (  # type: ignore
        get_report_blob, mark_report_paid, get_decrypted_report
    )

    entry = get_report_blob(audit_id)
    if not entry:
        raise HTTPException(status_code=404, detail="Report not found")

    if entry.get("paid"):
        # Already paid — return directly
        _, text, _ = get_decrypted_report(audit_id)
        return {"audit_id": audit_id, "status": "already_paid", "report_text": text}

    # Determine expected receiver
    try:
        from agents.agent_wallets import get_agent_address  # type: ignore
        reporting_wallet = get_agent_address("reporting_agent")
    except Exception:
        reporting_wallet = None

    bc = get_blockchain_client()
    expected_receiver = reporting_wallet or bc.address or ""

    # Verify on-chain (10 s timeout)
    try:
        from agents.x402_payments import verify_payment_on_chain  # type: ignore
        verified, err = verify_payment_on_chain(
            tx_id=body.tx_id,
            expected_receiver=expected_receiver,
            expected_amount_micro=int(REPORT_PAYMENT_ALGO * 1_000_000),
            max_wait_sec=10,
        )
    except Exception as exc:
        verified, err = False, str(exc)

    if not verified:
        # Return pending state — never block the caller forever
        return {
            "audit_id": audit_id,
            "status": "pending",
            "tx_id": body.tx_id,
            "message": f"Payment not yet confirmed: {err}. Retry in ~5 seconds.",
        }

    # Payment confirmed — decrypt and deliver
    mark_report_paid(audit_id, body.tx_id)
    _, text, decrypt_err = get_decrypted_report(audit_id)

    if not text:
        raise HTTPException(status_code=500, detail=f"Decryption failed: {decrypt_err}")

    return {
        "audit_id": audit_id,
        "status": "confirmed",
        "tx_id": body.tx_id,
        "report_text": text,
    }


# ══════════════════════════════════════════════════════════════════════════
# Part 5 — Revenue Dashboard Endpoints
# ══════════════════════════════════════════════════════════════════════════

@app.get("/api/revenue")
def get_revenue_dashboard() -> Dict[str, Any]:
    """
    GET /api/revenue

    Aggregates data from data/agent_payments.json for the Revenue tab.
    Returns:
      - total ALGO earned per agent
      - total audits paid for
      - total reports sold
      - agent wallet balances
      - last 10 X402 payments with TX IDs
    """
    payments_path = DATA_DIR / "agent_payments.json"
    payments: list = []
    if payments_path.exists():
        try:
            raw = payments_path.read_text(encoding="utf-8").strip() or "[]"
            payments = json.loads(raw)
        except Exception:
            payments = []

    # ── Aggregate by agent ────────────────────────────────────────────────
    agent_earnings: Dict[str, float] = {}
    total_audits_paid = 0
    total_reports_sold = 0

    for p in payments:
        agent = p.get("agent", "unknown")
        direction = p.get("direction", "outgoing")
        amount = p.get("amount_algo", 0.0)
        service = p.get("service", "")
        status = p.get("status", "")

        if direction == "incoming" and status == "confirmed":
            agent_earnings[agent] = agent_earnings.get(agent, 0.0) + amount
            if service == "audit":
                total_audits_paid += 1
            elif service == "report_access":
                total_reports_sold += 1

    # ── Agent wallet balances ─────────────────────────────────────────────
    wallets_path = DATA_DIR / "agent_wallets.json"
    wallet_addresses: Dict[str, str] = {}
    if wallets_path.exists():
        try:
            raw = wallets_path.read_text(encoding="utf-8").strip() or "{}"
            wd = json.loads(raw)
            for name, info in wd.items():
                if isinstance(info, dict):
                    wallet_addresses[name] = info.get("address", "")
        except Exception:
            pass

    # Try to fetch live balances (best effort — may fail if algod is down)
    agent_balances: Dict[str, Any] = {}
    for name, addr in wallet_addresses.items():
        agent_balances[name] = {
            "address": addr,
            "balance_algo": None,  # populated below
        }

    try:
        from algosdk.v2client import algod as _algod  # type: ignore
        _server = os.getenv("ALGOD_SERVER", "https://testnet-api.algonode.cloud")
        _token = os.getenv("ALGOD_TOKEN", "")
        if _token:
            _ac = _algod.AlgodClient(_token, _server)
        else:
            _ac = _algod.AlgodClient("", _server, headers={"User-Agent": "CfoE"})

        for name, info in agent_balances.items():
            addr = info.get("address", "")
            if addr:
                try:
                    acct_info = _ac.account_info(addr)
                    bal = acct_info.get("amount", 0) / 1_000_000
                    agent_balances[name]["balance_algo"] = round(bal, 6)
                except Exception:
                    pass
    except Exception:
        pass

    # ── Last 10 payments ──────────────────────────────────────────────────
    recent_payments = payments[:10]
    for p in recent_payments:
        tx = p.get("tx_id")
        if tx:
            p["explorer_url"] = f"https://testnet.algoexplorer.io/tx/{tx}"

    return {
        "total_algo_earned": round(sum(agent_earnings.values()), 6),
        "earnings_by_agent": {k: round(v, 6) for k, v in agent_earnings.items()},
        "total_audits_paid": total_audits_paid,
        "total_reports_sold": total_reports_sold,
        "agent_balances": agent_balances,
        "recent_payments": recent_payments,
        "payment_count": len(payments),
    }


@app.get("/api/agent-wallets")
def get_agent_wallets() -> Dict[str, Any]:
    """
    GET /api/agent-wallets

    Returns agent wallet addresses and live balances (polled by Revenue dashboard).
    """
    try:
        from agents.agent_wallets import initialize_agent_wallets  # type: ignore
        info = initialize_agent_wallets()
        return {"agents": info, "count": len(info)}
    except Exception as exc:
        raise HTTPException(status_code=500, detail=str(exc))


# Store active WebSocket connections
active_websockets: List[WebSocket] = []

@app.websocket("/ws/logs")
async def websocket_logs(websocket: WebSocket):
    await websocket.accept()
    active_websockets.append(websocket)
    try:
        while True:
            # Check for new log messages
            if not log_queue.empty():
                log_msg = log_queue.get()
                await websocket.send_json(log_msg)
            await asyncio.sleep(0.1)
    except WebSocketDisconnect:
        if websocket in active_websockets:
            active_websockets.remove(websocket)
    except Exception:
        if websocket in active_websockets:
            active_websockets.remove(websocket)


# ── Simulator Integration ─────────────────────────────────────────────

from simulator.simulator import (
    state as sim_state,
    manager as sim_manager,
    _build_snapshot,
    _current_shift,
    _compute_esg_score,
    PROCESSES,
    VIOLATION_TYPES,
    TICK_INTERVAL,
    MAX_HISTORY
)


@app.websocket("/ws/simulator")
async def ws_simulator(websocket: WebSocket) -> None:
    await sim_manager.connect(websocket)
    await websocket.send_json(_build_snapshot())
    try:
        while True:
            await websocket.receive_text()
    except WebSocketDisconnect:
        sim_manager.disconnect(websocket)


@app.post("/simulation/start")
async def start_simulation() -> Dict[str, Any]:
    sim_state.running = True
    return {"status": "running"}


@app.post("/simulation/stop")
async def stop_simulation() -> Dict[str, Any]:
    sim_state.running = False
    return {"status": "stopped"}


@app.post("/simulation/reset")
async def reset_simulation() -> Dict[str, Any]:
    sim_state.reset()
    await sim_manager.broadcast(_build_snapshot())
    return {"status": "reset"}


@app.post("/simulation/trigger-spike")
async def trigger_spike() -> Dict[str, Any]:
    sim_state.spike_active = True
    sim_state.spike_remaining = 6
    sim_state.spike_multiplier = round(__import__("random").uniform(1.4, 2.1), 2)
    return {"status": "spike_triggered", "multiplier": sim_state.spike_multiplier, "ticks": 6}


@app.post("/simulation/trigger-violation")
async def trigger_violation() -> Dict[str, Any]:
    entry = {"type": "EPA_FINE", "severity": "HIGH", "description": "EPA emission limit fine", "fine": 250_000, "id": f"VIO-{uuid4().hex[:8].upper()}", "timestamp": datetime.now(timezone.utc).isoformat()}
    sim_state.active_violations.append(entry)
    if len(sim_state.active_violations) > 10:
        sim_state.active_violations = sim_state.active_violations[-10:]
    sim_state.cumulative_violations += 3
    sim_state.esg_score = _compute_esg_score()
    await sim_manager.broadcast(_build_snapshot())
    return {"status": "violation_injected", "violation": entry}


@app.get("/simulation/snapshot")
async def snapshot() -> Dict[str, Any]:
    return _build_snapshot()


@app.post("/audit/run")
async def run_audit_from_simulator() -> Dict[str, Any]:
    snap = _build_snapshot()
    payload = {"supplier_name": snap["supplier_name"], "emissions": snap["estimated_annual_co2"], "violations": snap["cumulative_violations"], "notes": f"Simulator audit", "sector": "default"}
    try:
        result = run_audit(AuditRequest(**payload))
        result["download_links"] = export_audit_files(result)
        if result.get("human_approval_required", False):
            pending = load_pending()
            pending.insert(0, result)
            save_pending(pending)
        else:
            history = load_history()
            history.insert(0, result)
            save_history(history[:500])
        sim_state.last_audit_result = result
        await sim_manager.broadcast({"type": "audit_result", "result": result})
        return result
    except Exception as exc:
        error_msg = {"type": "audit_error", "error": str(exc)}
        await sim_manager.broadcast(error_msg)
        return error_msg


async def _simulation_loop() -> None:
    import random
    while True:
        if not sim_state.running:
            await asyncio.sleep(0.5)
            continue
        sim_state.tick_count += 1
        shift_name, shift_mult = _current_shift()
        tick_co2 = 0.0
        for proc in PROCESSES:
            noise = random.gauss(1.0, 0.06)
            daily_rate = proc["base_co2"] * shift_mult * noise
            if sim_state.spike_active:
                daily_rate *= sim_state.spike_multiplier
            per_tick = daily_rate / (86400 / TICK_INTERVAL)
            sim_state.process_emissions[proc["name"]] = daily_rate
            tick_co2 += per_tick
        sim_state.total_co2_today += tick_co2
        if sim_state.spike_active:
            sim_state.spike_remaining -= 1
            if sim_state.spike_remaining <= 0:
                sim_state.spike_active = False
                sim_state.spike_multiplier = 1.0
        else:
            if random.random() < 0.08:
                sim_state.spike_active = True
                sim_state.spike_remaining = random.randint(3, 8)
                sim_state.spike_multiplier = round(random.uniform(1.4, 2.1), 2)
        if random.random() < 0.05:
            viol = random.choice(VIOLATION_TYPES)
            entry = {**viol, "id": f"VIO-{uuid4().hex[:8].upper()}", "timestamp": datetime.now(timezone.utc).isoformat()}
            sim_state.active_violations.append(entry)
            if len(sim_state.active_violations) > 10:
                sim_state.active_violations = sim_state.active_violations[-10:]
            sim_state.cumulative_violations += 1
        sim_state.esg_score = _compute_esg_score()
        snap = _build_snapshot()
        sim_state.history.append(snap)
        if len(sim_state.history) > MAX_HISTORY:
            sim_state.history = sim_state.history[-MAX_HISTORY:]
        try:
            await sim_manager.broadcast(snap)
        except Exception:
            pass
        await asyncio.sleep(TICK_INTERVAL)


@app.on_event("startup")
async def startup_simulator() -> None:
    asyncio.create_task(_simulation_loop())
